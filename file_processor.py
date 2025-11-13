# file_processor.py (★ v3m/v3f のエラー対策・ハイブリッド版 ★)
 
import os
import re
from docx import Document
import pdfplumber
from openpyxl import load_workbook
import unicodedata
import pandas as pd
import io # 👈 ★ 1. io (BytesIO) をインポート
import base64 # 👈 ★ 2. base64 (デコード用) をインポート
 
import logging
logging.getLogger("pdfminer").setLevel(logging.ERROR)
 
# ----------------------------------------------------
# ユーティリティ関数: 各ファイル形式のテキスト化
# (★ file_path を file_path_or_bytes に変更 ★)
# ----------------------------------------------------
 
def extract_text_from_xlsx(file_path_or_bytes) -> str:
    """
    Excelファイル（パスまたはbytes）からテキストを抽出
    (★ 10万行の空行を無視する高速版 ★)
    """
    full_text = []
    wb = None
    try:
        wb = load_workbook(file_path_or_bytes, read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # --- ▼▼▼ ★★★ 修正点 ★★★ ▼▼▼
            # .iter_rows() は10万行を律儀に読んでしまうため、
            # .values (データがあるセルだけを読む) に変更する
            print(f"DEBUG (file_processor): '{sheet_name}' を .values モードで高速読み込み中...")
            for row_tuple in ws.values: 
                # row_tuple は (val1, val2, None, val4, ...) のようなタプル
                # 1. row_tuple が (None, None, None) のような
                #    「完全に空の行」であれば、ここで処理を打ち切る
                #    (※ただし、ws.valuesがどこまで返すかはopenpyxlの実装によるため、
                #     row_textのチェックも併用する)
                if not any(row_tuple):
                    continue # 完全に空の行はスキップ
 
                # 2. 値があるセルだけを結合する
                row_text = " ".join([str(cell_value) for cell_value in row_tuple if cell_value is not None])
                if row_text.strip(): # 結合した結果、意味のある文字列がある行のみ追加
                    full_text.append(row_text.strip())
            print(f"DEBUG (file_processor): '{sheet_name}' の読み込み完了。")
            # --- ▲▲▲ ★★★ 修正ここまで ★★★ ▲▲▲
 
        return "\n".join(full_text)
    except Exception as e:
        print(f"DEBUG (file_processor): extract_text_from_xlsx エラー: {e}") 
        return "" 
    finally:
        if wb:
            try:
                wb.close()
            except Exception as close_error:
                 pass
 
def extract_text_from_pdf(file_path_or_bytes) -> str:
    """PDFファイル（パスまたはbytes）からテキストを抽出"""
    text = ""
    try:
        # ★ file_path_or_bytes がパスかbytesかを自動判別 ★
        with pdfplumber.open(file_path_or_bytes) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text(x_tolerance=1, keep_blank_chars=False) or ""
                text += extracted + "\n"
        if not text.strip():
            return ""
 
        return text.strip()
    except Exception as e:
        return ""
 
 
def extract_text_from_docx(file_path_or_bytes) -> str:
    """Docxファイル（パスまたはbytes）からテキストを抽出"""
    full_text = []
    try:
        # ★ file_path_or_bytes がパスかbytesかを自動判別 ★
        document = Document(file_path_or_bytes)
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)
        for i, table in enumerate(document.tables):
            for row in table.rows:
                row_text = " ".join([cell.text.replace('\n', ' ').strip() for cell in row.cells])
                full_text.append(row_text)
        return "\n".join(filter(None, full_text))
    except Exception as e:
        return ""
 
# --- ▼▼▼ ★★★ 3. get_attachment_text を修正 ★★★ ▼▼▼
def get_attachment_text(filename: str, temp_file_path: str = None, content_bytes_base64: str = None) -> str:
    """
    添付ファイルのテキストを抽出する。
    win32com (temp_file_path) と Graph API/高速版 (content_bytes_base64) の両方に対応。
    """
    file_extension = os.path.splitext(filename)[1].lower()
    # ★ 処理対象（ファイルパス or バイナリ）を決定
    source_data = None
    if content_bytes_base64:
        try:
            # Base64 データをデコードしてバイナリにする
            source_data_bytes = base64.b64decode(content_bytes_base64)
            # バイナリデータをメモリ上で扱えるようにする
            source_data = io.BytesIO(source_data_bytes)
        except Exception as e:
            print(f"ERROR: Base64デコード失敗 (ファイル: {filename}): {e}")
            return ""
    elif temp_file_path:
        source_data = temp_file_path
    else:
        return "" # ソースがない
 
    # ★ 決定したソースを各抽出関数に渡す
    if file_extension in ['.xlsx', '.xls']:
        raw_text = extract_text_from_xlsx(source_data)
    elif file_extension == '.pdf':
        raw_text = extract_text_from_pdf(source_data)
    elif file_extension == '.docx':
        raw_text = extract_text_from_docx(source_data)
    else:
        return "" # 非対応なら空文字列
    # --- ▲▲▲▲▲▲▲▲▲▲▲▲▲▲▲▲▲▲▲▲▲▲▲ ---
 
    # --- 抽出後の最終クリーンアップ ---
    if not raw_text or pd.isna(raw_text):
         return ""
 
    cleaned_text = str(raw_text).strip()
    if cleaned_text.startswith("[ERROR:") or cleaned_text.startswith("[WARN:"):
         return ""
 
    try:
        cleaned_text = unicodedata.normalize('NFKC', cleaned_text)
        control_chars = ''.join(map(chr, list(range(0, 9)) + list(range(11, 13)) + list(range(14, 32)) + [127]))
        cleaned_text = re.sub(f'[{control_chars}\u200B\uFEFF]', '', cleaned_text)
        cleaned_text = re.sub(r'[\s\u3000]+', ' ', cleaned_text)
        cleaned_text = re.sub(r'(\s*\n\s*)+', '\n', cleaned_text)
        cleaned_text = cleaned_text.strip()
    except Exception as e:
         return str(raw_text).strip()
 
    return cleaned_text